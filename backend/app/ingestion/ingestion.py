import os
from io import BytesIO
from typing import List, Dict, Any
import fitz
import docx
from sqlalchemy import func
from sqlalchemy.orm import Session
from backend.app.db import Document, DocumentChunk, refresh_fts_index
from backend.app.llm.llm_client import transcribe_audio

def extract_text_from_pdf(file_content: bytes, doc_id: int) -> List[Dict[str, Any]]:
    """Extracts text, page numbers, and embedded XREF images from a PDF."""
    doc = fitz.open(stream=file_content, filetype="pdf")
    pages_data = []
    
    # Ensure static directory exists
    static_dir = os.path.join("backend", "app", "static", "images")
    os.makedirs(static_dir, exist_ok=True)

    for i, page in enumerate(doc):
        page_num = i + 1
        text = page.get_text()
        
        # Image extraction (Deep XREF isolation)
        image_paths = []
        image_list = page.get_images(full=True)
        for img_idx, img in enumerate(image_list):
            try:
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                img_ext = base_image["ext"]
                
                img_name = f"doc_{doc_id}_pg_{page_num}_img_{img_idx}.{img_ext}"
                img_path = os.path.join(static_dir, img_name)
                with open(img_path, "wb") as f:
                    f.write(image_bytes)
                image_paths.append(f"/static/images/{img_name}")
            except Exception as e:
                print(f"Error extracting image {img_idx} from page {page_num}: {e}")
                
        image_path_str = ",".join(image_paths) if image_paths else None

        if text.strip() or image_path_str:
            pages_data.append({
                "page_number": page_num,
                "content": text.strip() if text else "",
                "image_path": image_path_str
            })
    return pages_data

def extract_text_from_docx(file_content: bytes) -> List[Dict[str, Any]]:
    """Extracts text from a DOCX."""
    doc = docx.Document(BytesIO(file_content))
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Since DOCX doesn't have reliable page numbers without extra libs, 
    # we treat it as a single page or chunk it differently.
    return [{
        "page_number": 1,
        "content": "\n".join(full_text).strip()
    }]

def extract_text_from_txt(file_content: bytes) -> List[Dict[str, Any]]:
    """Extracts text from a TXT file."""
    text = file_content.decode("utf-8")
    return [{
        "page_number": 1,
        "content": text.strip()
    }]

def extract_text_from_doc(file_content: bytes) -> List[Dict[str, Any]]:
    """Extracts text from a .doc file with cross-platform fallbacks."""
    import subprocess
    import tempfile
    import sys
    
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tf:
        tf.write(file_content)
        temp_path = tf.name
    
    try:
        if sys.platform == "darwin":
            # macOS natively supports textutil securely
            result = subprocess.run(
                ["textutil", "-convert", "txt", "-stdout", temp_path],
                capture_output=True,
                text=True
            )
            if result.returncode == 0:
                return [{"page_number": 1, "content": result.stdout.strip()}]
                
        elif sys.platform == "win32":
            # Windows fallback execution logic
            try:
                # If pywin32 is installed, try native Word COM object translation
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                wb = word.Documents.Open(os.path.abspath(temp_path))
                doc_text = wb.Content.Text
                wb.Close()
                word.Quit()
                return [{"page_number": 1, "content": doc_text.strip()}]
            except ImportError:
                # Rudimentary fallback binary scraping
                print("Warning: pywin32 not found. Falling back to raw binary string extraction for .doc on Windows.")
                import string
                import re
                binary_text = file_content.decode('utf-8', errors='ignore')
                printable = set(string.printable)
                extracted = "".join(filter(lambda x: x in printable, binary_text))
                cleaned_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', '', extracted)
                cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
                return [{"page_number": 1, "content": cleaned_text}]
                
        elif sys.platform.startswith("linux"):
            # Linux typically uses antiword
            try:
                result = subprocess.run(["antiword", temp_path], capture_output=True, text=True)
                if result.returncode == 0:
                    return [{"page_number": 1, "content": result.stdout.strip()}]
            except FileNotFoundError:
                print("Warning: 'antiword' command not found on Linux. Skipping .doc text extraction.")
                
        return []
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

def extract_text_from_audio(file_content: bytes, file_ext: str) -> List[Dict[str, Any]]:
    """Transcribes audio content using OpenAI Whisper."""
    import tempfile
    
    with tempfile.NamedTemporaryFile(suffix=file_ext, delete=False) as tf:
        tf.write(file_content)
        temp_path = tf.name
    
    try:
        text = transcribe_audio(temp_path)
        return [{
            "page_number": 1,
            "content": text.strip()
        }]
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

def chunk_text(text: str, chunk_size=400, overlap=50) -> List[str]:
    """Chunks text into segments of approximately chunk_size words."""
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i : i + chunk_size])
        if chunk:
            chunks.append(chunk)
    return chunks


def _next_id(db: Session, model) -> int:
    current_max = db.query(func.max(model.id)).scalar()
    return 1 if current_max is None else int(current_max) + 1

def process_file(db: Session, file_name: str, file_content: bytes, file_ext: str):
    """Main ingestion pipeline."""
    # 1. Insert into documents table
    db_doc = Document(id=_next_id(db, Document), file_name=file_name)
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)
    next_chunk_id = _next_id(db, DocumentChunk)
    
    # 2. Extract text based on extension
    if file_ext == ".pdf":
        pages = extract_text_from_pdf(file_content, db_doc.id)
    elif file_ext == ".docx":
        pages = extract_text_from_docx(file_content)
    elif file_ext == ".doc":
        pages = extract_text_from_doc(file_content)
    elif file_ext == ".txt":
        pages = extract_text_from_txt(file_content)
    elif file_ext in [".mp3", ".wav", ".m4a", ".webm"]:
        pages = extract_text_from_audio(file_content, file_ext)
    else:
        return None

    # 3. Chunk and store
    for page in pages:
        content_text = page["content"]
        image_path = page.get("image_path")
        
        doc_header = f"File: {file_name} - Page {page['page_number']}"
        # If there's an image but no text, create a placeholder chunk to ensure the image is indexed
        if not content_text.strip() and image_path:
            db_chunk = DocumentChunk(
                id=next_chunk_id,
                document_id=db_doc.id,
                page_number=page["page_number"],
                content=f"[{doc_header} | Document Image Data]", # Placeholder for retrieval
                image_path=image_path
            )
            db.add(db_chunk)
            next_chunk_id += 1
            continue

        chunks = chunk_text(content_text)
        for i, chunk in enumerate(chunks):
            if not chunk.strip():
                continue
            db_chunk = DocumentChunk(
                id=next_chunk_id,
                document_id=db_doc.id,
                page_number=page["page_number"],
                content=f"[{doc_header}]\n{chunk}",
                image_path=image_path if i == 0 else None # Attaching image to first chunk only
            )
            db.add(db_chunk)
            next_chunk_id += 1
    
    # Force associate image to the first chunk if multiple chunks exist
    # This prevents the image from being duplicated across all chunks of the same page
    db.commit()
    refresh_fts_index()
    return db_doc.id
